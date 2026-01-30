import os
import json
import pandas as pd
import re
from datetime import datetime
from tkinter import filedialog, Tk
from docx import Document
from registro_fallos import registrar_fallo, limpiar_registro, mostrar_registro, LOG_FILE
from main import (
    obtener_rutas,
    insertar_imagen_con_transparencia,
    APPDATA_DIR,
    extraer_codigos_pdf,
    insertar_imagenes_en_pdf_placeholder,
)
from main import normalizar_cadena_alnum_mayus
from plantillaPDF import cargar_tabla_relacion

INDEX_FILE = os.path.join(APPDATA_DIR, "index_indice.json")
IMG_EXTS = [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif"]

# Debug logger for pegado indice
DEBUG_INDEX_LOG = os.path.join(APPDATA_DIR, "pegado_indice_debug.log")
def _log_index(msg: str):
    try:
        with open(DEBUG_INDEX_LOG, 'a', encoding='utf-8') as lf:
            lf.write(f"[{datetime.now().isoformat()}] {msg}\n")
    except Exception:
        try:
            print("[pegado_indice]", msg)
        except Exception:
            pass


def seleccionar_excel():
    Tk().withdraw()
    return filedialog.askopenfilename(
        title="Selecciona el archivo Excel para construir el índice",
        filetypes=[("Excel Files", "*.xlsx *.xlsm *.xls *.xlsb")]
    )


def normalizar_codigo(codigo):
    """Normaliza un código que puede venir como float/int/str.
    - Convierte NaN/None a "".
    - Si es float entero (ej. 28013578.0) devuelve "28013578".
    - Elimina sufijos ".0" en strings y quita espacios alrededor.
    - Devuelve string vacío si el resultado está vacío o es 'nan'.
    """
    try:
        import pandas as _pd
        if _pd.isna(codigo):
            return ""
    except Exception:
        pass

    if codigo is None:
        return ""

    # Enteros y floats
    if isinstance(codigo, float):
        # 28013578.0 -> '28013578'
        if codigo.is_integer():
            return str(int(codigo))
        # floats not integer: strip trailing zeros
        s = format(codigo, 'g')
        return s

    if isinstance(codigo, int):
        return str(codigo)

    s = str(codigo).strip()
    # Eliminar sufijo .0 que pandas a veces pone en strings
    if s.endswith('.0'):
        s = s[:-2]
    if s.lower() == 'nan':
        return ""
    return s


# Caché simple de listados de directorio para evitar os.listdir repetidos
_listdir_cache = {}
def _cached_listdir(path):
    try:
        key = os.path.normcase(os.path.abspath(path))
    except Exception:
        key = path
    if key in _listdir_cache:
        return _listdir_cache[key]
    try:
        items = os.listdir(path)
    except Exception:
        items = []
    _listdir_cache[key] = items
    return items


def construir_indice_desde_excel(ruta_excel):
    ext = os.path.splitext(ruta_excel)[1].lower()
    engine = "pyxlsb" if ext == ".xlsb" else None
    # Intentar leer forzando todas las columnas a strings para evitar floats
    try:
        df = pd.read_excel(ruta_excel, sheet_name="CONCENTRADO", engine=engine, dtype=str)
    except TypeError:
        # Algunos engines (pyxlsb) pueden no aceptar dtype; hacer fallback sin dtype
        try:
            df = pd.read_excel(ruta_excel, sheet_name="CONCENTRADO", engine=engine)
        except Exception as e:
            msg = str(e).lower()
            if ext == ".xlsb" and ("pyxlsb" in msg or "missing optional dependency" in msg):
                raise Exception("Missing optional dependency 'pyxlsb'.") from e
            raise
    except Exception as e:
        msg = str(e).lower()
        if ext == ".xlsb" and ("pyxlsb" in msg or "missing optional dependency" in msg):
            raise Exception("Missing optional dependency 'pyxlsb'.") from e
        raise

    indice = {}

    # Cargar tabla de relación para filtrar códigos válidos
    try:
        df_rel = cargar_tabla_relacion()
        valid_codes = set()
        for col in ("CODIGO","CODIGOS","CODE","SKU","CLAVE"):
            if col in df_rel.columns:
                # Normalizar valores numéricos leídos por pandas (evitar '28007960.0')
                def _cell_to_str(v):
                    try:
                        import pandas as _pd
                        if _pd.isna(v):
                            return ""
                    except Exception:
                        pass
                    # Si viene como float entero, convertir a int para quitar .0
                    try:
                        if isinstance(v, float) and v.is_integer():
                            return str(int(v))
                    except Exception:
                        pass
                    return str(v).strip()

                for v in df_rel[col].tolist():
                    s = _cell_to_str(v)
                    if not s:
                        continue
                    valid_codes.add(normalizar_cadena_alnum_mayus(s))
                break
    except Exception:
        valid_codes = None

    # Detectar columnas de código y destino por encabezado si es posible
    code_col = None
    dest_col = None
    try:
        cols = list(df.columns)
        # Normalizar encabezados para detección
        def _norm_col(c):
            return re.sub(r"[^A-Za-z0-9]", "", str(c or "")).upper()

        norm = {c: _norm_col(c) for c in cols}
        # Buscar columna de código
        for c, nc in norm.items():
            if any(k in nc for k in ("CODIGO", "CODIGOS", "CODE", "SKU", "EAN", "UPC", "ESTILO")):
                code_col = c
                break

        # Buscar columna de destino/asignación (columna B esperada: ASIGNACIÓN/ASIG/DESTINO)
        for c, nc in norm.items():
            if any(k in nc for k in ("ASIG", "ASIGN", "ASIGNACION", "DESTINO", "NOMBRE", "EVIDENCIA")):
                dest_col = c
                break

        # Log columns chosen
        _log_index(f"construir_indice: detected cols code_col={code_col} dest_col={dest_col} cols={cols}")
    except Exception:
        code_col = None
        dest_col = None

    # Helper: normalizar valores de código/destino
    def _cell_str_from_row(val):
        return normalizar_codigo(val)

    for _, row in df.iterrows():
        try:
            if code_col is not None and dest_col is not None:
                codigo = _cell_str_from_row(row.get(code_col, ""))
                destino = _cell_str_from_row(row.get(dest_col, ""))
            else:
                # Fallback antiguo: columna A -> código, columna B -> destino
                codigo = _cell_str_from_row(row.iloc[0])
                destino = _cell_str_from_row(row.iloc[1])
        except Exception:
            continue

        if not codigo or not destino or codigo.lower() == "nan" or destino.lower() == "nan":
            continue

        if "código" in codigo.lower() or "sku" in codigo.lower():
            continue

        canon = normalizar_cadena_alnum_mayus(codigo)
        # Si disponemos de la tabla de relación, solo añadimos códigos que estén en ella
        if valid_codes is not None and canon not in valid_codes:
            continue

        indice[canon] = destino
        # También añadir la variante literal que pandas/otros puedan usar (p.ej. '28013578.0')
        try:
            raw_key = str(codigo).strip()
            if raw_key and raw_key != canon:
                indice[raw_key] = destino
        except Exception:
            pass

    with open(INDEX_FILE, "w", encoding="utf-8") as f:
        json.dump(indice, f, ensure_ascii=False, indent=4)
    _log_index(f"Indice construido: {len(indice)} entries; index_file={INDEX_FILE}; sample_keys={list(indice.keys())[:10]}")

    return indice


def extraer_codigos_tabla(doc):
    codigos = []
    for tabla in doc.tables:
        if not tabla.rows:
            continue

        encabezados = [c.text.strip().upper() for c in tabla.rows[0].cells]
        if any("CODIGO" in h.replace("Ó", "O") or "SKU" in h or "CLAVE" in h for h in encabezados):
            idx = 0
            for i, h in enumerate(encabezados):
                h_norm = h.replace("Ó", "O")
                if "CODIGO" in h_norm or "SKU" in h_norm or "CLAVE" in h_norm:
                    idx = i
                    break

            for fila in tabla.rows[1:]:
                texto = (fila.cells[idx].text or "").strip()
                if not texto:
                    continue

                # Normalizar a sólo caracteres alfanuméricos para la clave
                canon = "".join(ch for ch in texto if ch.isalnum())
                if not canon:
                    continue

                codigos.append(texto)
                _log_index(f"extraer_codigos_tabla: encontrado -> original='{texto}' canon='{canon}'")

            break
    return codigos


def buscar_destino(ruta_base, destino):
    destino = destino.strip()

    base, ext = os.path.splitext(destino)

    if ext.lower() in IMG_EXTS:
        for archivo in _cached_listdir(ruta_base):
            if archivo.lower() == destino.lower():
                return "imagen", os.path.join(ruta_base, archivo)

    nombre_base = base if ext.lower() in IMG_EXTS else destino
    # Buscar coincidencias de nombre base, incluyendo variantes tipo "1234(2)", "1234-2", "1234_2"
    # Primer intento: buscar coincidencias directas en el directorio raíz
    matches = []
    nb = nombre_base.strip().lower()
    for archivo in _cached_listdir(ruta_base):
        archivo_base, archivo_ext = os.path.splitext(archivo)
        if archivo_ext.lower() not in IMG_EXTS:
            continue
        ab = archivo_base.strip().lower()
        if ab == nb:
            matches.append(os.path.join(ruta_base, archivo))
            continue
        try:
            m = re.match(rf"^{re.escape(nb)}(?:\s*\(\d+\)|[-_]\d+)$", ab, flags=re.IGNORECASE)
            if m:
                matches.append(os.path.join(ruta_base, archivo))
        except Exception:
            if ab.startswith(nb):
                rem = ab[len(nb):].strip()
                if rem.startswith('(') and rem.endswith(')') and rem[1:-1].isdigit():
                    matches.append(os.path.join(ruta_base, archivo))
                elif (rem.startswith('-') or rem.startswith('_')) and rem[1:].isdigit():
                    matches.append(os.path.join(ruta_base, archivo))

    if matches:
        if len(matches) == 1:
            return "imagen", matches[0]
        return "imagen", matches

    # Segundo intento: buscar recursivamente (busca carpetas y archivos con el nombre)
    try:
        for root, dirs, files in os.walk(ruta_base):
            # carpetas con el nombre exacto -> devolver carpeta
            for d in dirs:
                if d.strip().lower() == nb:
                    return "carpeta", os.path.join(root, d)
            # archivos con nombre base o nombre exacto
            for f in files:
                f_base, f_ext = os.path.splitext(f)
                if f_ext.lower() not in IMG_EXTS:
                    continue
                if f.strip().lower() == destino.strip().lower():
                    return "imagen", os.path.join(root, f)
                if f_base.strip().lower() == nb:
                    matches.append(os.path.join(root, f))
        if matches:
            if len(matches) == 1:
                return "imagen", matches[0]
            return "imagen", matches
    except Exception:
        pass

    carpeta_buscada = nombre_base
    for item in _cached_listdir(ruta_base):
        if os.path.isdir(os.path.join(ruta_base, item)) and item.lower() == carpeta_buscada.lower():
            return "carpeta", os.path.join(ruta_base, item)

    return None, None


def procesar_doc_con_indice_docx(ruta_doc, ruta_imagenes, indice):
    doc = Document(ruta_doc)
    codigos = extraer_codigos_tabla(doc)

    fallo_registrado = False
    imagenes_insertadas = 0

    # Recoger problemas detectados durante el procesamiento para reportar al final
    doc_issues = []

    if not codigos:
        registrar_fallo(os.path.basename(ruta_doc), reason="no_codes", details={"ruta_doc": ruta_doc})
        return

    _log_index(f"Procesando DOCX: {ruta_doc}; codigos_found={len(codigos)}; ruta_imagenes={ruta_imagenes}")
    for p in doc.paragraphs:
        txt = (p.text or "")
        # Accept case-variants like ${IMAGEN} or ${imagen} (allow spaces inside braces)
        if re.search(r"\$\{\s*imagen\s*\}", txt, flags=re.IGNORECASE):
            # Limpiar el párrafo de forma segura (python-docx no tiene `clear()` pública)
            try:
                for r in list(p.runs):
                    try:
                        p._element.remove(r._element)
                    except Exception:
                        # Si falla la eliminación directa, intentar borrar el texto
                        try:
                            r.text = ""
                        except Exception:
                            pass
            except Exception:
                try:
                    p.text = ""
                except Exception:
                    pass
            _log_index(f"placeholder encontrado en paragraph: {ruta_doc}")
            run = p.add_run()

            for codigo in codigos:
                codigo_norm = normalizar_codigo(codigo)
                canon = normalizar_cadena_alnum_mayus(codigo_norm)
                if canon not in indice and codigo_norm not in indice:
                    doc_issues.append({"type": "code_not_in_index", "codigo": codigo})
                    continue

                destino = indice[canon]
                _log_index(f"Codigo {codigo!r} -> canon={canon} -> destino {destino}")
                tipo, ruta = buscar_destino(ruta_imagenes, destino)
                _log_index(f"buscar_destino -> tipo={tipo} ruta={ruta}")

                if tipo == "imagen":
                    # `ruta` puede ser una lista (varios archivos con misma base)
                    if isinstance(ruta, (list, tuple)):
                        for rp in ruta:
                            insertar_imagen_con_transparencia(run, rp)
                            imagenes_insertadas += 1
                    else:
                        insertar_imagen_con_transparencia(run, ruta)
                        imagenes_insertadas += 1

                elif tipo == "carpeta":
                    archivos_en_carpeta = [a for a in _cached_listdir(ruta) if os.path.splitext(a)[1].lower() in IMG_EXTS]
                    if not archivos_en_carpeta:
                        doc_issues.append({"type": "carpeta_vacia", "carpeta": ruta, "codigo": codigo})
                    else:
                        for archivo in archivos_en_carpeta:
                            insertar_imagen_con_transparencia(run, os.path.join(ruta, archivo))
                            imagenes_insertadas += 1

                else:
                    doc_issues.append({"type": "destino_no_encontrado", "codigo": codigo, "destino": destino})

            break

    if imagenes_insertadas == 0:
        # Registrar un único fallo por documento con las razones detectadas
        detalles = {
            "codigos": codigos,
            "issues": doc_issues,
            "ruta_imagenes": ruta_imagenes,
        }
        registrar_fallo(os.path.basename(ruta_doc), reason="no_images_inserted", details=detalles)
        fallo_registrado = True

    doc.save(ruta_doc)
    print(f"Documento actualizado: {ruta_doc}")


def procesar_doc_con_indice_pdf(ruta_doc, ruta_imagenes, indice):
    codigos = extraer_codigos_pdf(ruta_doc)

    fallo_registrado = False
    imagenes_insertadas = 0

    # Recoger problemas detectados durante el procesamiento
    doc_issues = []

    if not codigos:
        registrar_fallo(os.path.basename(ruta_doc), reason="no_codes", details={"ruta_doc": ruta_doc})
        return

    rutas_imagenes = []

    _log_index(f"Procesando PDF: {ruta_doc}; codigos_found={len(codigos)}; ruta_imagenes={ruta_imagenes}")
    for codigo in codigos:
        codigo_norm = normalizar_codigo(codigo)
        canon = normalizar_cadena_alnum_mayus(codigo_norm)
        if canon not in indice and codigo_norm not in indice:
            doc_issues.append({"type": "code_not_in_index", "codigo": codigo})
            continue

        # Preferir destino por canon, si no existe usar la clave literal
        destino = indice.get(canon) or indice.get(codigo_norm)
        _log_index(f"Codigo {codigo!r} -> canon={canon} -> destino {destino}")
        tipo, ruta = buscar_destino(ruta_imagenes, destino)
        _log_index(f"buscar_destino -> tipo={tipo} ruta={ruta}")

        if tipo == "imagen":
            # ruta puede ser lista
            if isinstance(ruta, (list, tuple)):
                for rp in ruta:
                    rutas_imagenes.append(rp)
                    imagenes_insertadas += 1
            else:
                rutas_imagenes.append(ruta)
                imagenes_insertadas += 1

        elif tipo == "carpeta":
            archivos_en_carpeta = [a for a in _cached_listdir(ruta) if os.path.splitext(a)[1].lower() in IMG_EXTS]
            if not archivos_en_carpeta:
                doc_issues.append({"type": "carpeta_vacia", "carpeta": ruta, "codigo": codigo})
            else:
                for archivo in archivos_en_carpeta:
                    rutas_imagenes.append(os.path.join(ruta, archivo))
                    imagenes_insertadas += 1

        else:
            doc_issues.append({"type": "destino_no_encontrado", "codigo": codigo, "destino": destino})

    if imagenes_insertadas == 0:
        detalles = {
            "codigos": codigos,
            "issues": doc_issues,
            "ruta_imagenes": ruta_imagenes,
        }
        registrar_fallo(os.path.basename(ruta_doc), reason="no_images_inserted", details=detalles)
        fallo_registrado = True
        return

    exito = insertar_imagenes_en_pdf_placeholder(ruta_doc, rutas_imagenes)
    if not exito and not fallo_registrado:
        registrar_fallo(os.path.basename(ruta_doc), reason="insert_failed", details={"ruta_doc": ruta_doc, "rutas_imagenes_len": len(rutas_imagenes)})


def procesar_doc_con_indice(ruta_doc, ruta_imagenes, indice):
    ext = os.path.splitext(ruta_doc)[1].lower()
    if ext == ".docx":
        procesar_doc_con_indice_docx(ruta_doc, ruta_imagenes, indice)
    elif ext == ".pdf":
        procesar_doc_con_indice_pdf(ruta_doc, ruta_imagenes, indice)


def procesar_indice():
    limpiar_registro()

    ruta_docs, ruta_imgs = obtener_rutas()
    if not ruta_docs or not ruta_imgs:
        return

    excel = seleccionar_excel()
    if not excel:
        raise Exception("No se seleccionó un archivo Excel para el modo Pegado por Índice.")

    print("Construyendo índice desde Excel...")
    indice = construir_indice_desde_excel(excel)
    print("Índice generado correctamente.")

    archivos = [
        f for f in _cached_listdir(ruta_docs)
        if (f.endswith(".docx") or f.endswith(".pdf")) and not f.startswith("~$")
    ]

    for archivo in archivos:
        procesar_doc_con_indice(os.path.join(ruta_docs, archivo), ruta_imgs, indice)

    mostrar_registro()

    if os.path.exists(LOG_FILE):
        os.startfile(LOG_FILE)


        