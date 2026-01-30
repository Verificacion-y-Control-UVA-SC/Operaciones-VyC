import os
import json
import sys
import re
import unicodedata
from tkinter import Tk, filedialog

from docx import Document
from docx.shared import Inches
from PIL import Image

# ============================================================
# INTENTAMOS IMPORTAR REGISTRO_DE_FALLOS (SIN CICLOS)
# ============================================================
try:
    from registro_fallos import registrar_fallo, limpiar_registro, mostrar_registro
except Exception as e:
    print(f"Error al importar registro_fallos: {e}")
    registrar_fallo = None
    limpiar_registro = None
    mostrar_registro = None

# ============================================================
# PYMUPDF (PDF)
# ============================================================
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except Exception as e:
    HAS_PYMUPDF = False
    print("PyMuPDF (fitz) no está instalado. El modo PDF no estará disponible:", e)

# ============================================================
# STDOUT UTF-8
# ============================================================
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

# ============================================================
# APPDATA / CONFIG / LOG
# ============================================================
APPDATA_DIR = os.path.join(os.getenv("APPDATA"), "ImagenesVC")
os.makedirs(APPDATA_DIR, exist_ok=True)

CONFIG_FILE = os.path.join(APPDATA_DIR, "config.json")
LOG_FILE = os.path.join(APPDATA_DIR, "documentos_sin_imagenes.txt")

# ============================================================
# CONSTANTES DE FILTRADO / EXTENSIONES
# ============================================================
FORBIDDEN_TOKENS = {
    "TOTAL", "CANTIDAD", "FACTURA", "MARCA", "DESCRIPCION", "DESCRIPCIÓN",
    "FECHA", "CONTRATO", "PRESENTACION", "PRESENTACIÓN", "SISTEMA",
    "ALEATORIO", "DICTAMEN", "PRODUCTO", "RELACION", "RELACIÓN",
    "MODELO", "ORIGEN", "CHINA", "MALASIA", "ALEMANIA", "RUMANIA",
    "ITALIA", "BRASIL", "DIMENSIONES", "CONTENIDO", "ETIQUETA", "CONTEN",
}
IMG_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}

# Dimensiones usadas para escalado DOCX
H_MAX_W_CM = 4.36
H_MAX_H_CM = 6.37
V_MAX_W_CM = 8.13
V_MAX_H_CM = 4.84


# ============================================================
# CONFIGURACIÓN
# ============================================================
def guardar_config(data):
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)


def cargar_config():
    if not os.path.exists(CONFIG_FILE):
        return {}
    try:
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def seleccionar_carpeta(titulo):
    Tk().withdraw()
    carpeta = filedialog.askdirectory(title=titulo)
    return carpeta.replace("\\", "/").strip() if carpeta else None


def obtener_rutas():
    """
    Devuelve (ruta_docs, ruta_imagenes) usando config en APPDATA.
    Si alguna falta, la pide al usuario y la guarda.
    """
    cfg = cargar_config()
    ruta_docs = cfg.get("ruta_docs", "")
    ruta_imgs = cfg.get("ruta_imagenes", "")

    if not os.path.isdir(ruta_docs):
        ruta_docs = seleccionar_carpeta("Selecciona la carpeta de documentos .docx/.pdf")
        if not ruta_docs:
            return None, None
        cfg["ruta_docs"] = ruta_docs

    if not os.path.isdir(ruta_imgs):
        ruta_imgs = seleccionar_carpeta("Selecciona la carpeta de imágenes")
        if not ruta_imgs:
            return None, None
        cfg["ruta_imagenes"] = ruta_imgs

    guardar_config(cfg)
    return ruta_docs, ruta_imgs


# ============================================================
# NORMALIZACIÓN / UTILIDADES
# ============================================================
def _sin_acentos(s):
    return "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")


def normalizar_cadena_alnum_mayus(s):
    return re.sub(r"[^A-Za-z0-9]", "", s or "").upper()


def contiene_digito(s):
    return any(c.isdigit() for c in (s or ""))


def norm_path_key(path):
    return os.path.normcase(os.path.normpath(path or ""))


# ============================================================
# INDEXADO DE IMÁGENES (CARPETA ÚNICA)
# ============================================================
def indexar_imagenes(carpeta_imagenes):
    """
    Construye un índice de las imágenes en una carpeta:
      - name
      - base (sin extensión)
      - base_norm (solo alfanumérico mayúsculas)
      - path
    """
    index = []
    def _core_base(name):
        # Eliminar sufijos comunes como ' (2)', '-2', '_2' al final del nombre
        import re
        core = re.sub(r"[\s\-_]*\(\s*\d+\s*\)$", "", name)
        core = re.sub(r"[\s\-_]+\d+$", "", core)
        return core

    for nombre in os.listdir(carpeta_imagenes):
        base, ext = os.path.splitext(nombre)
        if ext.lower() not in IMG_EXTS:
            continue
        core = _core_base(base)
        index.append({
            "name": nombre,
            "base": base,
            "base_core": core,
            "ext": ext,
            "base_norm": normalizar_cadena_alnum_mayus(base),
            "base_core_norm": normalizar_cadena_alnum_mayus(core),
            "path": os.path.join(carpeta_imagenes, nombre),
        })
    return index


def buscar_imagen_index(index, codigo_canonico, usadas_paths, usadas_bases):
    """
    Busca una imagen en el índice para un código dado usando:
      1) coincidencia exacta con base_norm
      2) coincidencia parcial con heurística simple
    Evita repetir paths o bases ya usadas.
    """
    code = normalizar_cadena_alnum_mayus(codigo_canonico)
    if not code:
        return None

    # Coincidencia exacta (tanto con base_norm como con base_core_norm)
    exactos = [
        it for it in index
        if (it.get("base_norm") == code or it.get("base_core_norm") == code)
        and norm_path_key(it["path"]) not in usadas_paths
        and (it.get("base_norm") not in usadas_bases and it.get("base_core_norm") not in usadas_bases)
    ]
    if exactos:
        return exactos[0]["path"]

    # Coincidencias parciales
    parciales = [
        it for it in index
        if ((code in (it.get("base_norm") or "") or (it.get("base_norm") or "") in code)
            or (code in (it.get("base_core_norm") or "") or (it.get("base_core_norm") or "") in code))
        and norm_path_key(it["path"]) not in usadas_paths
        and (it.get("base_norm") not in usadas_bases and it.get("base_core_norm") not in usadas_bases)
    ]
    if not parciales:
        return None

    def score(it):
        bn = it["base_norm"]
        starts = bn.startswith(code)
        ends = bn.endswith(code)
        delta = abs(len(bn) - len(code))
        return (0 if starts or ends else 1, delta, bn)

    parciales.sort(key=score)
    return parciales[0]["path"]


# ============================================================
# EXTRACCIÓN DE CÓDIGOS DESDE DOCX
# ============================================================
def extraer_codigos(doc):
    """
    Extrae códigos desde tablas en un DOCX.
    Usa patrones generales y un patrón especial tipo Bosch.
    Devuelve lista sin duplicados.
    """
    codigos = []
    patron_general = re.compile(r"[A-Za-z0-9][A-Za-z0-9.\-]{4,}", re.IGNORECASE)
    patron_bosch = re.compile(r"(?:No\.?\s*)?(\d(?:\s?\d){8,12})")

    for tabla in doc.tables:
        if not tabla.rows:
            continue

        idx_codigo = None

        # Buscamos índice de columna de código
        for r in range(min(3, len(tabla.rows))):
            for i, celda in enumerate(tabla.rows[r].cells):
                t_norm = _sin_acentos(celda.text or "").upper()
                if ("CODIGO" in t_norm or "SKU" in t_norm or "CLAVE" in t_norm):
                    idx_codigo = i
                    break
            if idx_codigo is not None:
                break

        columnas = [idx_codigo] if idx_codigo is not None else range(len(tabla.rows[0].cells))

        # Recorremos filas de datos
        for fila in tabla.rows[1:]:
            for j in columnas:
                texto = (fila.cells[j].text or "").strip()

                # Bosch numérico
                for m in patron_bosch.findall(texto):
                    num = re.sub(r"\s+", "", m)
                    if 8 <= len(num) <= 13:
                        codigos.append(num)

                # Patrones generales
                for m in patron_general.findall(texto):
                    canon = normalizar_cadena_alnum_mayus(m)
                    if canon and contiene_digito(canon) and canon not in FORBIDDEN_TOKENS:
                        codigos.append(canon)

    # Unificar
    return list(dict.fromkeys(codigos))


# ============================================================
# EXTRACCIÓN DE CÓDIGOS DESDE PDF
# ============================================================
def extraer_codigos_pdf(ruta_pdf):
    """
    Extrae códigos desde la TABLA del dictamen PDF.
    Busca encabezados como: MARCA | CÓDIGO | FACTURA | CANTIDAD
    y devuelve el contenido exacto de la columna CÓDIGO.

    Esta función solo afecta a PDFs. El pegado para DOCX sigue usando
    extraer_codigos() / extraer_codigos_tabla() y no se ve modificado.
    """
    if not HAS_PYMUPDF:
        print("PyMuPDF no disponible; no se pueden leer códigos desde PDF.")
        return []

    codigos = []

    try:
        doc = fitz.open(ruta_pdf)
    except Exception as e:
        print(f"No se pudo abrir PDF {ruta_pdf}: {e}")
        return []

    for page in doc:
        try:
            tablas = page.find_tables()
        except Exception as e:
            print(f"No se pudieron detectar tablas en {ruta_pdf}: {e}")
            continue

        for tabla in tablas:
            try:
                df = tabla.to_pandas()
            except Exception as e:
                print(f"No se pudo convertir tabla a pandas en {ruta_pdf}: {e}")
                continue

            # Normalizar encabezados
            encabezados = [
                str(h).strip().upper().replace("Ó", "O")
                for h in df.columns
            ]

            # Buscar índice de la columna de CÓDIGO
            idx_codigo = None
            for i, h in enumerate(encabezados):
                if "CODIGO" in h or "SKU" in h or "CLAVE" in h:
                    idx_codigo = i
                    break

            if idx_codigo is None:
                continue

            # Extraer códigos de esa columna
            for valor in df.iloc[:, idx_codigo].tolist():
                if valor is None:
                    continue

                texto = str(valor).strip()
                if not texto or texto.lower() == "nan":
                    continue

                canon = normalizar_cadena_alnum_mayus(texto)

                # Validar que parezca un código real (letras y números)
                if len(canon) < 6:
                    continue
                if not any(c.isdigit() for c in canon):
                    continue
                if not any(c.isalpha() for c in canon):
                    continue

                codigos.append(texto)

    doc.close()

    # Eliminar duplicados manteniendo el orden
    codigos_unicos = []
    vistos = set()
    for c in codigos:
        if c not in vistos:
            vistos.add(c)
            codigos_unicos.append(c)

    return codigos_unicos

# ============================================================
# INSERCIÓN DE IMAGEN (DOCX)
# ============================================================
def insertar_imagen_con_transparencia(run, img_path):
    """
    Inserta una imagen en un run, escalándola para respetar
    las dimensiones máximas configuradas y evitando tapar encabezados.
    """
    try:
        with Image.open(img_path) as img:
            w_px, h_px = img.size

        # Conversión a pulgadas asumiendo 96 dpi
        w_in = w_px / 96.0
        h_in = h_px / 96.0

        if w_px >= h_px:
            max_w_in = H_MAX_W_CM / 2.54
            max_h_in = H_MAX_H_CM / 2.54
        else:
            max_w_in = V_MAX_W_CM / 2.54
            max_h_in = V_MAX_H_CM / 2.54

        scale = min(max_w_in / w_in, max_h_in / h_in, 1)
        new_w_in = w_in * scale
        new_h_in = h_in * scale

        if new_w_in <= max_w_in:
            run.add_picture(img_path, width=Inches(new_w_in))
        else:
            run.add_picture(img_path, height=Inches(new_h_in))

        run.add_text(" ")
    except Exception as e:
        print(f"Error al insertar {img_path}: {e}")


# ============================================================
# INSERCIÓN EN PDF (PLACEHOLDER)
# ============================================================
def insertar_imagenes_en_pdf_placeholder(ruta_pdf, rutas_imagenes, placeholder="${imagen}"):
    """
    Inserta una o varias imágenes en un PDF usando un marcador de texto
    (por defecto ${imagen}).

    - Busca el placeholder en el texto.
    - Borra el texto del marcador (redacción blanca).
    - Coloca las imágenes en cuadrícula debajo del marcador, dentro
      de los márgenes de la página.

    Devuelve True si se intentó insertar en el PDF y se guardó correctamente.
    """
    if not HAS_PYMUPDF:
        print("PyMuPDF no disponible; no se pueden modificar PDFs.")
        return False

    if not rutas_imagenes:
        print(f"No hay imágenes para insertar en {ruta_pdf}.")
        return False

    # Eliminar duplicados y rutas inexistentes antes de abrir el PDF
    try:
        import os as _os
        seen = set()
        uniq = []
        missing = []
        for p in rutas_imagenes:
            # admitir que cada entrada pueda ser un dict con clave 'imagen_path' u otras
            candidate = None
            try:
                if isinstance(p, dict):
                    # Priorizar llaves comunes
                    candidate = p.get('imagen_path') or p.get('path') or p.get('ruta') or p.get('imagen')
                else:
                    candidate = p
            except Exception:
                candidate = p

            if not candidate:
                missing.append(p)
                continue

            try:
                path_norm = _os.path.normcase(_os.path.normpath(str(candidate)))
            except Exception:
                path_norm = str(candidate)

            if not _os.path.exists(path_norm):
                missing.append(candidate)
                continue

            # deduplicar por ruta normalizada (permitir múltiples archivos con mismo contenido)
            if path_norm in seen:
                continue
            seen.add(path_norm)
            uniq.append(path_norm)

        if missing:
            try:
                print(f"Advertencia: {len(missing)} imágenes no existen y serán omitidas. Ejemplo: {missing[:3]}")
            except Exception:
                pass

        rutas_imagenes = uniq
        if not rutas_imagenes:
            print(f"No quedan imágenes válidas tras deduplicación/validación para {ruta_pdf}.")
            return False
    except Exception:
        pass
    try:
        doc = fitz.open(ruta_pdf)
    except Exception as e:
        print(f"No se pudo abrir PDF para inserción de imágenes: {ruta_pdf}: {e}")
        return False

    page_target = None
    marca = None

    # Buscar la primera ocurrencia del placeholder en todo el documento.
    # Hacer búsqueda tolerante a mayúsculas/minúsculas probando variantes.
    variants = [placeholder]
    try:
        variants.append(placeholder.upper())
        variants.append(placeholder.lower())
    except Exception:
        pass

    for page in doc:
        found = False
        for v in variants:
            try:
                rects = page.search_for(v)
            except Exception:
                rects = []
            if rects:
                page_target = page
                marca = rects[0]
                found = True
                break
        if found:
            break

    if page_target is None or marca is None:
        print(f"No se encontró el marcador {placeholder} en {ruta_pdf}")
        doc.close()
        return False

    page_rect = page_target.rect

    # Borramos el texto del marcador usando redacción
    try:
        page_target.add_redact_annot(marca, fill=(1, 1, 1))
        page_target.apply_redactions()
    except Exception as e:
        print(f"No se pudo aplicar redacción sobre marcador en {ruta_pdf}: {e}")

    # Parámetros de cuadrícula en puntos (1 pulgada = 72 puntos)
    max_w_in = H_MAX_W_CM / 2.54
    max_h_in = H_MAX_H_CM / 2.54
    max_w_pt = max_w_in * 72
    max_h_pt = max_h_in * 72

    espacio_x = 8   # puntos
    espacio_y = 8
    por_fila = 3

    # Posición inicial: debajo del marcador, pero forzando márgenes razonables
    margen_izq = page_rect.x0 + 36  # 0.5"
    margen_sup = page_rect.y0 + 36
    margen_inf = page_rect.y1 - 36

    start_x = max(marca.x0, margen_izq)
    start_y = max(marca.y1 + 12, margen_sup)  # 12 puntos debajo del marcador
    if start_y + max_h_pt > margen_inf:
        # Si se sale por abajo, movemos hacia arriba
        start_y = max(margen_sup, margen_inf - max_h_pt * 2)

    # Inserción de imágenes
    # Mantener conjunto de rutas insertadas para evitar insertar exactamente
    # la misma ruta más de una vez (pero permitir distintos paths aunque
    # tengan contenido idéntico, ya que el usuario puede querer ese comportamiento).
    inserted_paths = set()
    # Además mantener hashes insertados para evitar pegar la misma imagen
    # visual más de una vez aunque venga desde paths distintos.
    inserted_hashes = set()
    import hashlib
    # Funciones de hashing: por defecto usar hash normalizado de imagen
    def _file_md5(p):
        try:
            h = hashlib.md5()
            with open(p, 'rb') as fh:
                for chunk in iter(lambda: fh.read(8192), b''):
                    h.update(chunk)
            return h.hexdigest()
        except Exception:
            return None

    def _image_normalized_hash(p, size=(64, 64)):
        try:
            # Usar PIL para abrir, convertir a RGB, redimensionar y hash de bytes
            with Image.open(p) as im:
                im = im.convert('RGB')
                im = im.resize(size, resample=Image.LANCZOS)
                data = im.tobytes()
            return hashlib.md5(data).hexdigest()
        except Exception:
            return None

    # Leer configuración para decidir modo de deduplicación por contenido
    try:
        cfg_local = cargar_config() if 'cargar_config' in globals() else {}
    except Exception:
        cfg_local = {}
    DEDUPE_CONTENT = bool(cfg_local.get('dedupe_by_content', False))
    try:
        print(f"Iniciando inserción en {ruta_pdf}: {len(rutas_imagenes)} imágenes; dedupe_by_content={DEDUPE_CONTENT}. Ejemplo: {rutas_imagenes[:3]}")
    except Exception:
        pass
    for idx, img_path in enumerate(rutas_imagenes):
        fila = idx // por_fila
        col = idx % por_fila

        x0 = start_x + col * (max_w_pt + espacio_x)
        y0 = start_y + fila * (max_h_pt + espacio_y)
        x1 = x0 + max_w_pt
        y1 = y0 + max_h_pt

        # Si nos salimos por el ancho de página, intentamos mover a la siguiente fila
        if x1 > page_rect.x1 - 20:
            fila += 1
            col = 0
            x0 = start_x + col * (max_w_pt + espacio_x)
            y0 = start_y + fila * (max_h_pt + espacio_y)
            x1 = x0 + max_w_pt
            y1 = y0 + max_h_pt

        # Si ahora se sale por el alto de la página, lo registramos y saltamos
        if y1 > page_rect.y1 - 20:
            print(f"Advertencia: no hay espacio para insertar imagen en {ruta_pdf}: {img_path}")
            continue

        rect = fitz.Rect(x0, y0, x1, y1)

        try:
            try:
                pnorm = os.path.normcase(os.path.normpath(img_path))
            except Exception:
                pnorm = img_path
            if pnorm in inserted_paths:
                print(f"Omitida inserción duplicada por ruta: {img_path}")
                continue

            # calcular hash y omitir si ya se insertó una imagen idéntica
            img_hash = None
            if DEDUPE_CONTENT:
                img_hash = _image_normalized_hash(pnorm)
                # si falla el método normalizado, caer al hash de archivo
                if img_hash is None:
                    img_hash = _file_md5(pnorm)
            else:
                img_hash = _file_md5(pnorm)

            if img_hash is not None and img_hash in inserted_hashes:
                print(f"Omitida inserción duplicada por contenido (hash): {img_path}")
                # marcar ruta como vista para evitar procesarla otra vez
                inserted_paths.add(pnorm)
                continue

            inserted_paths.add(pnorm)
            if img_hash is not None:
                inserted_hashes.add(img_hash)

            page_target.insert_image(rect, filename=img_path, keep_proportion=True)
            print(f"Imagen insertada en PDF {ruta_pdf}: {img_path}")
        except Exception as e:
            print(f"Error al insertar imagen en PDF {ruta_pdf}: {e}")

    # Guardar cambios usando archivo temporal para evitar el error
    # "save to original must be incremental"
    temp_path = ruta_pdf + ".tmp"

    try:
        # Guardado completo (no incremental) en archivo temporal
        doc.save(temp_path, incremental=False, encryption=fitz.PDF_ENCRYPT_KEEP)
        doc.close()

        # Reemplazar el archivo original por el temporal
        if os.path.exists(ruta_pdf):
            os.remove(ruta_pdf)
        os.rename(temp_path, ruta_pdf)

        print(f"PDF actualizado: {ruta_pdf}")
        return True

    except Exception as e:
        print(f"Error al guardar PDF {ruta_pdf}: {e}")
        try:
            doc.close()
        except Exception:
            pass

        # Si quedó el temporal colgando, lo eliminamos
        try:
            if os.path.exists(temp_path):
                os.remove(temp_path)
        except Exception:
            pass

        return False

# ============================================================
# PROCESO PRINCIPAL (ORQUESTADOR DE MODOS)
# ============================================================
def procesar_lote():
    """
    Fuerza reconstrucción de rutas y archivos en cada ejecución,
    sin reutilizar estado previo de sesiones anteriores.
    """

    # Siempre cargar config desde archivo
    cfg = cargar_config()
    modo = (cfg.get("modo_pegado", "simple") or "").lower()

    # Siempre obtener rutas de documentos e imágenes en cada corrida
    ruta_docs, ruta_imgs = obtener_rutas()
    if not ruta_docs or not ruta_imgs:
        print("No hay rutas configuradas.")
        return

    # Mensaje de diagnóstico (útil para verificar que está ejecutando)
    print(f"Procesando lote en modo: {modo}")
    print(f"Ruta de documentos: {ruta_docs}")
    print(f"Ruta de imágenes: {ruta_imgs}")

    # Desvío por modo
    if modo == "indice":
        from pegado_indice import procesar_indice
        procesar_indice()
        return

    if modo == "carpetas":
        from pegado_carpetas import procesar_carpetas
        procesar_carpetas()
        return

    # Si no es índice ni carpetas → modo simple
    from pegado_simple import procesar_simple
    procesar_simple()

# ============================================================
# EJECUCIÓN DIRECTA (OPCIONAL)
# ============================================================
if __name__ == "__main__":
    procesar_lote()
