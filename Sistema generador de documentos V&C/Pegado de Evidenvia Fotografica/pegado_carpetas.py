import os
from docx import Document
from main import (
    obtener_rutas,
    insertar_imagen_con_transparencia,
    extraer_codigos,
    normalizar_cadena_alnum_mayus,
    extraer_codigos_pdf,
    insertar_imagenes_en_pdf_placeholder,
)
from plantillaPDF import cargar_tabla_relacion
from registro_fallos import registrar_fallo, limpiar_registro, mostrar_registro, LOG_FILE

IMG_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}


def construir_indice_carpetas(ruta_imgs):
    """
    Crea un índice de carpetas:
        clave normalizada (solo letras/números, mayúsculas) -> [rutas de carpeta]
    Esto permite que códigos como 'KI1545138' encuentren carpetas llamadas 'KI154-5138'.
    """
    indice = {}

    try:
        entries = os.listdir(ruta_imgs)
    except Exception:
        entries = []

    for nombre in entries:
        ruta = os.path.join(ruta_imgs, nombre)
        if not os.path.isdir(ruta):
            continue

        clave = normalizar_cadena_alnum_mayus(nombre)
        if not clave:
            continue

        if clave not in indice:
            indice[clave] = []
        indice[clave].append(ruta)

    print(f"Índice de carpetas construido con {len(indice)} claves.")
    return indice


def procesar_carpetas():
    limpiar_registro()

    ruta_docs, ruta_imgs = obtener_rutas()
    if not ruta_docs or not ruta_imgs:
        return

    carpetas_index = construir_indice_carpetas(ruta_imgs)

    try:
        docs_entries = os.listdir(ruta_docs)
    except Exception:
        docs_entries = []
    archivos = [
        f for f in docs_entries
        if (f.endswith(".docx") or f.endswith(".pdf")) and not f.startswith("~$")
    ]

    for archivo in archivos:
        ruta_doc = os.path.join(ruta_docs, archivo)
        ext = os.path.splitext(archivo)[1].lower()

        if ext == ".docx":
            print(f"Procesando documento (modo carpetas DOCX): {ruta_doc}")
            doc = Document(ruta_doc)

            imagen_insertada = False
            codigos = extraer_codigos(doc)

            # Filtrar contra tabla de relación
            try:
                df_rel = cargar_tabla_relacion()
                valid_codes = set()
                for col in ("CODIGO","CODIGOS","CODE","SKU","CLAVE"):
                    if col in df_rel.columns:
                        for v in df_rel[col].astype(str).fillna(""):
                            valid_codes.add(normalizar_cadena_alnum_mayus(v))
                        break
            except Exception:
                valid_codes = None

            if codigos:
                if valid_codes is not None:
                    codigos = [c for c in codigos if normalizar_cadena_alnum_mayus(c) in valid_codes]

            if not codigos:
                print("  No se encontraron códigos en el documento (o ninguno coincide con la tabla de relación).")
                registrar_fallo(archivo)
                continue

            for p in doc.paragraphs:
                text_lower = (p.text or "").lower()
                if "${imagen}" in text_lower or "${imagen}".upper() in (p.text or ""):
                    p.clear()
                    run = p.add_run()

                    for codigo in codigos:
                        clave = normalizar_cadena_alnum_mayus(codigo)
                        if not clave:
                            continue

                        carpetas = carpetas_index.get(clave, [])
                        if not carpetas:
                            print(f"  No se encontró carpeta para código '{codigo}' (clave '{clave}').")
                            continue

                        for carpeta_codigo in carpetas:
                            try:
                                files_in_code = os.listdir(carpeta_codigo)
                            except Exception:
                                files_in_code = []
                            for archivo_img in files_in_code:
                                ext_img = os.path.splitext(archivo_img)[1].lower()
                                if ext_img not in IMG_EXTS:
                                    continue

                                # Solo aceptar imágenes cuyo nombre base normalizado
                                # coincida exactamente con la clave del código.
                                base = os.path.splitext(archivo_img)[0]
                                if normalizar_cadena_alnum_mayus(base) != clave:
                                    # Ignorar imágenes que no coincidan exactamente
                                    continue

                                img_path = os.path.join(carpeta_codigo, archivo_img)
                                insertar_imagen_con_transparencia(run, img_path)
                                imagen_insertada = True
                                print(f"  Imagen insertada: {img_path}")

                    break

            if not imagen_insertada:
                registrar_fallo(archivo)

            doc.save(ruta_doc)
            print(f"Documento actualizado: {ruta_doc}")

        elif ext == ".pdf":
            print(f"Procesando documento (modo carpetas PDF): {ruta_doc}")

            imagen_insertada = False
            codigos = extraer_codigos_pdf(ruta_doc)

            # Filtrar contra tabla de relación
            try:
                df_rel = cargar_tabla_relacion()
                valid_codes = set()
                for col in ("CODIGO","CODIGOS","CODE","SKU","CLAVE"):
                    if col in df_rel.columns:
                        for v in df_rel[col].astype(str).fillna(""):
                            valid_codes.add(normalizar_cadena_alnum_mayus(v))
                        break
            except Exception:
                valid_codes = None

            if codigos:
                if valid_codes is not None:
                    codigos = [c for c in codigos if normalizar_cadena_alnum_mayus(c) in valid_codes]

            if not codigos:
                print("  No se encontraron códigos en el PDF (o ninguno coincide con la tabla de relación).")
                registrar_fallo(archivo)
                continue

            rutas_imagenes = []

            for codigo in codigos:
                clave = normalizar_cadena_alnum_mayus(codigo)
                if not clave:
                    continue

                carpetas = carpetas_index.get(clave, [])
                if not carpetas:
                    print(f"  No se encontró carpeta para código '{codigo}' (clave '{clave}').")
                    continue

                for carpeta_codigo in carpetas:
                    try:
                        files_in_code = os.listdir(carpeta_codigo)
                    except Exception:
                        files_in_code = []
                    for archivo_img in files_in_code:
                        ext_img = os.path.splitext(archivo_img)[1].lower()
                        if ext_img not in IMG_EXTS:
                            continue

                        # Solo aceptar imágenes cuyo nombre base normalizado
                        # coincida exactamente con la clave del código.
                        base = os.path.splitext(archivo_img)[0]
                        if normalizar_cadena_alnum_mayus(base) != clave:
                            continue

                        img_path = os.path.join(carpeta_codigo, archivo_img)
                        rutas_imagenes.append(img_path)
                        imagen_insertada = True
                        print(f"  Imagen detectada para PDF: {img_path}")

            if not rutas_imagenes:
                registrar_fallo(archivo)
                continue

            # Intentar insertar usando placeholder en minúsculas; si no funciona, intentar mayúsculas
            exito = insertar_imagenes_en_pdf_placeholder(ruta_doc, rutas_imagenes, placeholder="${imagen}")
            if not exito:
                exito = insertar_imagenes_en_pdf_placeholder(ruta_doc, rutas_imagenes, placeholder="${IMAGEN}")
            if not exito:
                registrar_fallo(archivo)

    mostrar_registro()
    if os.path.exists(LOG_FILE):
        os.startfile(LOG_FILE)